import logging
from functools import singledispatchmethod

from docx.shared import Parented, RGBColor

from .renderable import *
from .renderable import Renderable
from . import extended_markdown
from .renderable.table import Table
from .renderable.caption import Caption
from .renderable.formula import Formula
from .renderable.heading import Heading
from .renderable.list import List
from .renderable.toc import ToC


class RenderableFactory:
    @singledispatchmethod
    @staticmethod
    def create(marko_element: extended_markdown.BlockElement, parent: Parented) \
            -> Renderable:
        paragraph = Paragraph(parent)
        paragraph.add_run(f"{marko_element.get_type()} is not supported", color=RGBColor.from_string('ff0000'))
        logging.warning(f"{marko_element.get_type()} is not supported")
        return paragraph

    @staticmethod
    def _create_runs(paragraph: Paragraph, children, classes: list[type] = None):
        if not classes:
            classes = []
        for child in children:
            if isinstance(child, (extended_markdown.RawText, extended_markdown.Literal)):
                paragraph.add_run(child.children,
                                  extended_markdown.StrongEmphasis in classes or None,
                                  extended_markdown.Emphasis in classes or None)
            elif isinstance(child, extended_markdown.CodeSpan):
                paragraph.add_run(child.children, is_italic=True)
            elif isinstance(child, extended_markdown.Image):
                paragraph.add_image(child.dest)
            elif isinstance(child, extended_markdown.LineBreak):
                pass  # ignore
            elif isinstance(child, (extended_markdown.Link, extended_markdown.Url)):
                paragraph.add_link(child.children[0].children,
                                   child.dest,
                                   extended_markdown.StrongEmphasis in classes or None,
                                   extended_markdown.Emphasis in classes or None)
            elif isinstance(child, (extended_markdown.Emphasis, extended_markdown.StrongEmphasis)):
                RenderableFactory._create_runs(paragraph, child.children, classes+[type(child)])
            else:
                paragraph.add_run(f" {child.get_type()} is not supported ", color=RGBColor.from_string("FF0000"))
                logging.warning(f"{child.get_type()} is not supported")

    @create.register
    @staticmethod
    def _(marko_paragraph: extended_markdown.Paragraph, parent: Parented):
        paragraph = Paragraph(parent)
        RenderableFactory._create_runs(paragraph, marko_paragraph.children)
        return paragraph

    @create.register
    @staticmethod
    def _(marko_heading: extended_markdown.Heading, parent: Parented):
        heading = Heading(parent, marko_heading.level, marko_heading.numbered)
        RenderableFactory._create_runs(heading, marko_heading.children)
        return heading

    @create.register
    @staticmethod
    def _(marko_code_block: extended_markdown.FencedCode | extended_markdown.CodeBlock, parent: Parented):
        listing = Listing(parent, marko_code_block.lang)
        listing.set_text(marko_code_block.children[0].children)
        return listing

    @create.register
    @staticmethod
    def _(marko_formula: extended_markdown.Formula, parent: Parented):
        formula = Formula(parent, marko_formula.formula)
        return formula

    @create.register
    @staticmethod
    def _(marko_caption: extended_markdown.Caption, parent: Parented):
        caption = Caption(parent, marko_caption.type, marko_caption.text)
        return caption

    @create.register
    @staticmethod
    def _(marko_list: extended_markdown.List, parent: Parented):
        list_ = List(parent, marko_list.ordered)

        def create_items_from_marko(marko_list_, level=1):
            for list_item in marko_list_.children:
                for child in list_item.children:
                    if isinstance(child, extended_markdown.List):
                        create_items_from_marko(child, level + 1)
                    elif isinstance(child, extended_markdown.Paragraph):
                        RenderableFactory._create_runs(
                            list_.add_item(level),
                            child.children
                        )

        create_items_from_marko(marko_list)

        return list_

    @create.register
    @staticmethod
    def _(marko_table: extended_markdown.Table, parent: Parented):
        table = Table(parent, len(marko_table.children), len(marko_table.children[0].children))

        for i, row in enumerate(marko_table.children):
            for j, cell in enumerate(row.children):
                RenderableFactory._create_runs(
                    table.add_paragraph_to_cell(i, j),
                    cell.children
                )

        return table

    @create.register
    @staticmethod
    def _(marko_toc: extended_markdown.TOC, parent: Parented):
        toc = ToC(parent)
        return toc
